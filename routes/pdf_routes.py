# ==========================================
# FILE: routes/pdf_routes.py (CORRECTED)
# ==========================================
"""
PDF Routes - File-based storage with Google Drive Integration
"""
from fastapi import APIRouter, File, UploadFile, HTTPException, Header
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from pathlib import Path
from datetime import datetime
from typing import List
from io import BytesIO
import uuid
import os
import json
import shutil
import subprocess

from config import settings
from schemas import DownloadRequest
from utils.pdf_processor import (
    extract_text_blocks, render_page_as_image, 
    save_thumbnail, apply_text_edits, apply_signatures
)
import fitz
from pydantic import BaseModel
from typing import List, Optional

# ✅ NEW IMPORTS FOR CONVERTERS
from PIL import Image
import img2pdf

router = APIRouter(prefix="/api", tags=["PDF"])

# In-memory storage for session data
SESSIONS = {}

class EditData(BaseModel):
    page: int
    bbox: List[float]
    old_text: str
    new_text: str
    fontSize: int
    color: str

class SignatureData(BaseModel):
    id: int
    page: int
    x: float
    y: float
    width: float
    height: float
    image: str

class PDFMergeItem(BaseModel):
    session_id: str
    edits: List[EditData] = []
    signatures: List[SignatureData] = []

class MergeEditedRequest(BaseModel):
    pdfs: List[PDFMergeItem]


@router.post("/upload")
async def upload_pdf(
    file: UploadFile = File(...),
    authorization: str = Header(None)
):
    """Upload and process PDF"""
    if not file.filename.endswith('.pdf'):
        raise HTTPException(400, "Only PDF files allowed")
    
    user_email = None
    if authorization:
        try:
            from routes.auth_routes import SESSIONS as AUTH_SESSIONS
            token = authorization.replace("Bearer ", "")
            if token in AUTH_SESSIONS:
                user_email = AUTH_SESSIONS[token].get("email")
        except ImportError:
            pass
    
    try:
        session_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        stored_filename = f"{timestamp}_{session_id[:8]}_{file.filename}"
        file_path = settings.UPLOAD_DIR / stored_filename
        
        print(f"📤 Upload - Session: {session_id}")
        if user_email:
            print(f"👤 User: {user_email}")
        
        file_data = await file.read()
        file_size = len(file_data)
        
        with open(file_path, 'wb') as f:
            f.write(file_data)
        
        doc = fitz.open(str(file_path))
        total_pages = len(doc)
        doc.close()
        
        all_pages = []
        for page_num in range(total_pages):
            page_img = render_page_as_image(str(file_path), page_num, 2.0)
            text_blocks = extract_text_blocks(str(file_path), page_num)
            
            thumb_filename = f"{session_id}_page_{page_num}.png"
            thumb_path = settings.THUMBNAIL_DIR / thumb_filename
            save_thumbnail(str(file_path), page_num, str(thumb_path))
            
            thumb_img = render_page_as_image(str(file_path), page_num, 0.3)
            
            all_pages.append({
                'page_num': page_num,
                'image': page_img['image'],
                'thumbnail': thumb_img['image'],
                'width': page_img['width'],
                'height': page_img['height'],
                'text_blocks': text_blocks
            })
        
        SESSIONS[session_id] = {
            'original_filename': file.filename,
            'stored_filename': stored_filename,
            'file_path': str(file_path),
            'file_size': file_size,
            'total_pages': total_pages,
            'uploaded_at': datetime.now().isoformat(),
            'pages': all_pages,
            'user_email': user_email
        }
        
        print(f"✅ Upload complete - {total_pages} pages")
        
        return JSONResponse({
            'success': True,
            'session_id': session_id,
            'pdf_id': session_id,
            'filename': file.filename,
            'total_pages': total_pages,
            'pages': all_pages
        })
        
    except Exception as e:
        print(f"❌ Upload error: {str(e)}")
        raise HTTPException(500, f"Upload failed: {str(e)}")


@router.post("/download")
async def download_pdf(request: DownloadRequest):
    """Download edited PDF and upload to Google Drive"""
    if len(request.edits) == 0 and len(request.signatures) == 0:
        raise HTTPException(400, "No changes to save")
    
    session_data = SESSIONS.get(request.session_id)
    if not session_data:
        raise HTTPException(404, "Session not found")
    
    try:
        original_path = session_data['file_path']
        
        if not Path(original_path).exists():
            raise HTTPException(404, "Original PDF file not found")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"edited_{timestamp}_{session_data['original_filename']}"
        output_path = settings.OUTPUT_DIR / output_filename
        
        temp_filename = f"temp_{timestamp}_{uuid.uuid4().hex[:8]}.pdf"
        temp_path = settings.OUTPUT_DIR / temp_filename
        
        current_input = original_path
        
        if len(request.edits) > 0:
            success = apply_text_edits(current_input, request.edits, str(temp_path))
            if not success:
                raise HTTPException(500, "Failed to apply text edits")
            current_input = str(temp_path)
        
        if len(request.signatures) > 0:
            if current_input == original_path:
                shutil.copy(original_path, temp_path)
                current_input = str(temp_path)
            
            success = apply_signatures(current_input, request.signatures, str(output_path))
            if not success:
                if Path(temp_path).exists():
                    os.remove(temp_path)
                raise HTTPException(500, "Failed to apply signatures")
        else:
            shutil.move(temp_path, output_path)
        
        if Path(temp_path).exists() and Path(temp_path) != Path(output_path):
            try:
                os.remove(temp_path)
            except:
                pass
        
        drive_link = None
        drive_file_id = None
        
        try:
            from routes.auth_routes import USERS
            
            user_email = session_data.get('user_email')
            
            if user_email and user_email in USERS:
                user_data = USERS[user_email]
                
                if user_data.get("drive_access_token"):
                    from services.drive_service import DriveService
                    
                    print(f"☁️ Uploading to Drive for: {user_email}")
                    
                    drive = DriveService(
                        user_data["drive_access_token"],
                        user_data.get("drive_refresh_token")
                    )
                    
                    drive_result = drive.upload_file(
                        str(output_path),
                        folder_id=user_data.get("drive_folder_id"),
                        file_name=output_filename
                    )
                    
                    if drive_result:
                        drive_link = drive_result['view_link']
                        drive_file_id = drive_result['id']
                        print(f"✅ Uploaded to Drive: {drive_link}")
                else:
                    print(f"⚠️ No Drive token for user: {user_email}")
            else:
                print(f"⚠️ No user linked to session")
                
        except Exception as e:
            print(f"⚠️ Drive upload failed (continuing): {e}")
            import traceback
            traceback.print_exc()
        
        return FileResponse(
            str(output_path),
            media_type='application/pdf',
            filename=output_filename,
            headers={
                "X-Drive-Link": drive_link or "not-uploaded",
                "X-Drive-File-ID": drive_file_id or ""
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Download error: {str(e)}")
        raise HTTPException(500, f"Download failed: {str(e)}")


@router.post("/merge-pdfs")
async def merge_pdfs(
    files: List[UploadFile] = File(...),
    authorization: str = Header(None)
):
    """Merge multiple PDF files into one"""
    if len(files) < 2:
        raise HTTPException(400, "At least 2 PDF files required for merging")
    
    for file in files:
        if not file.filename.endswith('.pdf'):
            raise HTTPException(400, f"File '{file.filename}' is not a PDF")
    
    try:
        from PyPDF2 import PdfMerger
        
        print(f"🔗 Starting merge of {len(files)} PDFs...")
        
        merger = PdfMerger()
        temp_files = []
        
        try:
            for idx, file in enumerate(files, 1):
                print(f"  📄 Processing file {idx}/{len(files)}: {file.filename}")
                
                content = await file.read()
                
                temp_filename = f"temp_merge_{uuid.uuid4().hex[:8]}_{file.filename}"
                temp_path = settings.UPLOAD_DIR / temp_filename
                temp_files.append(temp_path)
                
                with open(temp_path, 'wb') as f:
                    f.write(content)
                
                try:
                    merger.append(str(temp_path))
                    print(f"    ✓ Added successfully")
                except Exception as e:
                    print(f"    ✗ Failed to add: {str(e)}")
                    raise HTTPException(400, f"Failed to merge '{file.filename}': {str(e)}")
            
            print(f"  🔨 Writing merged PDF...")
            output = BytesIO()
            merger.write(output)
            merger.close()
            output.seek(0)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            merged_filename = f"merged_{timestamp}.pdf"
            
            output_size = output.getbuffer().nbytes
            print(f"✅ Merge complete: {merged_filename} ({output_size} bytes)")
            
            return StreamingResponse(
                output,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename={merged_filename}",
                    "X-Merged-Files": str(len(files)),
                    "X-Output-Size": str(output_size)
                }
            )
            
        except HTTPException:
            raise
        except Exception as e:
            print(f"❌ Merge error: {str(e)}")
            raise HTTPException(500, f"Merge failed: {str(e)}")
            
        finally:
            print(f"  🗑️ Cleaning up {len(temp_files)} temporary files...")
            for temp_file in temp_files:
                try:
                    if temp_file.exists():
                        os.remove(temp_file)
                        print(f"    ✓ Removed: {temp_file.name}")
                except Exception as e:
                    print(f"    ✗ Failed to remove {temp_file.name}: {str(e)}")
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Unexpected merge error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Merge operation failed: {str(e)}")


@router.post("/merge-edited-pdfs")
async def merge_edited_pdfs(
    request: MergeEditedRequest,
    authorization: str = Header(None)
):
    """Merge PDFs with their edits applied"""
    from PyPDF2 import PdfMerger
    
    pdfs_data = request.pdfs
    
    if len(pdfs_data) < 2:
        raise HTTPException(400, "At least 2 PDFs required")
    
    try:
        print(f"🔗 Merging {len(pdfs_data)} edited PDFs...")
        
        merger = PdfMerger()
        temp_files = []
        
        try:
            for idx, pdf_info in enumerate(pdfs_data, 1):
                session_id = pdf_info.session_id
                edits = [edit.dict() for edit in pdf_info.edits]
                signatures = [sig.dict() for sig in pdf_info.signatures]
                
                print(f"  📄 {idx}. Session: {session_id}")
                print(f"      Edits: {len(edits)}, Signatures: {len(signatures)}")
                
                session_data = SESSIONS.get(session_id)
                if not session_data:
                    raise HTTPException(404, f"Session {session_id} not found")
                
                original_path = session_data['file_path']
                
                if not Path(original_path).exists():
                    raise HTTPException(404, f"Original file not found: {original_path}")
                
                if len(edits) == 0 and len(signatures) == 0:
                    print(f"      Using original: {session_data['original_filename']}")
                    merger.append(original_path)
                    continue
                
                print(f"      Processing with changes: {session_data['original_filename']}")
                
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                temp_edited_path = settings.OUTPUT_DIR / f"temp_edited_{idx}_{timestamp}_{uuid.uuid4().hex[:8]}.pdf"
                temp_files.append(temp_edited_path)
                
                current_input = original_path
                
                if len(edits) > 0:
                    temp_edits_path = settings.OUTPUT_DIR / f"temp_edits_{idx}_{timestamp}_{uuid.uuid4().hex[:8]}.pdf"
                    temp_files.append(temp_edits_path)
                    
                    print(f"      Applying {len(edits)} text edits...")
                    success = apply_text_edits(current_input, edits, str(temp_edits_path))
                    if not success:
                        raise HTTPException(500, f"Failed to apply edits to PDF {idx}")
                    current_input = str(temp_edits_path)
                    print(f"      ✓ Text edits applied")
                
                if len(signatures) > 0:
                    print(f"      Applying {len(signatures)} signatures...")
                    if current_input == original_path:
                        shutil.copy(original_path, temp_edited_path)
                        current_input = str(temp_edited_path)
                    
                    success = apply_signatures(current_input, signatures, str(temp_edited_path))
                    if not success:
                        raise HTTPException(500, f"Failed to apply signatures to PDF {idx}")
                    print(f"      ✓ Signatures applied")
                else:
                    if current_input != original_path:
                        shutil.copy(current_input, temp_edited_path)
                
                merger.append(str(temp_edited_path))
                print(f"      ✓ Added to merge queue")
            
            print(f"  🔨 Writing merged PDF...")
            output = BytesIO()
            merger.write(output)
            merger.close()
            output.seek(0)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            merged_filename = f"merged_edited_{timestamp}.pdf"
            
            output_size = output.getbuffer().nbytes
            print(f"✅ Merge complete: {merged_filename} ({output_size} bytes)")
            
            return StreamingResponse(
                output,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename={merged_filename}",
                    "X-Merged-Files": str(len(pdfs_data)),
                    "X-Output-Size": str(output_size)
                }
            )
            
        finally:
            print(f"  🗑️ Cleaning up {len(temp_files)} temporary files...")
            for temp_file in temp_files:
                try:
                    if temp_file.exists():
                        os.remove(temp_file)
                        print(f"      ✓ Removed: {temp_file.name}")
                except Exception as e:
                    print(f"      ✗ Failed to remove {temp_file.name}: {str(e)}")
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Merge error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Merge failed: {str(e)}")


# ==========================================
# ✅ NEW: CONVERTER ENDPOINTS
# ==========================================

@router.post("/convert-images-to-pdf")
async def convert_images_to_pdf(
    images: List[UploadFile] = File(...),
    authorization: str = Header(None)
):
    """Convert multiple images to a single PDF"""
    if len(images) == 0:
        raise HTTPException(400, "No images provided")
    
    user_email = None
    if authorization:
        try:
            from routes.auth_routes import SESSIONS as AUTH_SESSIONS
            token = authorization.replace("Bearer ", "")
            if token in AUTH_SESSIONS:
                user_email = AUTH_SESSIONS[token].get("email")
        except:
            pass
    
    try:
        print(f"🖼️ Converting {len(images)} images to PDF...")
        temp_image_paths = []
        
        try:
            for idx, image in enumerate(images, 1):
                if not image.content_type.startswith('image/'):
                    raise HTTPException(400, f"'{image.filename}' is not an image")
                
                content = await image.read()
                temp_filename = f"temp_img_{uuid.uuid4().hex[:8]}_{image.filename}"
                temp_path = settings.UPLOAD_DIR / temp_filename
                temp_image_paths.append(temp_path)
                
                with open(temp_path, 'wb') as f:
                    f.write(content)
                
                # Validate and convert image
                img = Image.open(temp_path)
                if img.mode == 'RGBA':
                    rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                    rgb_img.paste(img, mask=img.split()[3])
                    rgb_img.save(temp_path)
                elif img.mode != 'RGB':
                    img.convert('RGB').save(temp_path)
            
            # Convert to PDF
            pdf_bytes = img2pdf.convert([str(path) for path in temp_image_paths])
            output = BytesIO(pdf_bytes)
            output.seek(0)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            pdf_filename = f"images_to_pdf_{timestamp}.pdf"
            
            print(f"✅ Conversion complete: {pdf_filename}")
            
            return StreamingResponse(
                output,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename={pdf_filename}",
                    "X-Image-Count": str(len(images))
                }
            )
            
        finally:
            for temp_file in temp_image_paths:
                try:
                    if temp_file.exists():
                        os.remove(temp_file)
                except:
                    pass
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Image to PDF error: {str(e)}")
        raise HTTPException(500, f"Conversion failed: {str(e)}")


@router.post("/convert-word-to-pdf")
async def convert_word_to_pdf(
    file: UploadFile = File(...),
    authorization: str = Header(None)
):
    """
    Convert Word document to PDF using python-docx + reportlab
    Works on Windows, Mac, Linux - No LibreOffice needed!
    """
    if not (file.filename.endswith('.doc') or file.filename.endswith('.docx')):
        raise HTTPException(400, "Only .doc and .docx files supported")
    
    try:
        # Try to import required libraries
        try:
            from docx import Document
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError as e:
            print(f"❌ Missing package: {str(e)}")
            raise HTTPException(
                500, 
                "Required packages not installed. Please run: pip install python-docx reportlab"
            )
        
        print(f"📝 Converting Word to PDF: {file.filename}")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        temp_word_filename = f"temp_word_{timestamp}_{uuid.uuid4().hex[:8]}_{file.filename}"
        temp_word_path = settings.UPLOAD_DIR / temp_word_filename
        
        temp_pdf_filename = f"temp_pdf_{timestamp}_{uuid.uuid4().hex[:8]}.pdf"
        temp_pdf_path = settings.OUTPUT_DIR / temp_pdf_filename
        
        try:
            # Save Word file
            content = await file.read()
            with open(temp_word_path, 'wb') as f:
                f.write(content)
            
            print(f"  💾 Saved Word file: {temp_word_path.name}")
            
            # Read Word document
            doc = Document(str(temp_word_path))
            
            # Create PDF with A4 page size
            pdf_doc = SimpleDocTemplate(
                str(temp_pdf_path),
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=36
            )
            
            # Container for PDF elements
            story = []
            styles = getSampleStyleSheet()
            
            # Define custom styles
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                leading=16,
                alignment=TA_LEFT,
                spaceAfter=10
            )
            
            heading1_style = ParagraphStyle(
                'CustomHeading1',
                parent=styles['Heading1'],
                fontSize=18,
                leading=22,
                textColor='#000000',
                spaceAfter=16,
                spaceBefore=12,
                fontName='Helvetica-Bold'
            )
            
            heading2_style = ParagraphStyle(
                'CustomHeading2',
                parent=styles['Heading2'],
                fontSize=14,
                leading=18,
                textColor='#000000',
                spaceAfter=12,
                spaceBefore=10,
                fontName='Helvetica-Bold'
            )
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=22,
                leading=26,
                textColor='#000000',
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            # Process document
            print(f"  📄 Processing {len(doc.paragraphs)} paragraphs...")
            
            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                
                if not text:
                    # Add small spacer for empty paragraphs
                    story.append(Spacer(1, 0.1 * inch))
                    continue
                
                # Escape special characters for PDF
                text = text.replace('&', '&amp;')
                text = text.replace('<', '&lt;')
                text = text.replace('>', '&gt;')
                
                # Determine style based on Word style
                style_name = para.style.name
                
                if style_name == 'Title':
                    p = Paragraph(text, title_style)
                elif style_name.startswith('Heading 1'):
                    p = Paragraph(text, heading1_style)
                elif style_name.startswith('Heading 2'):
                    p = Paragraph(text, heading2_style)
                elif style_name.startswith('Heading'):
                    p = Paragraph(text, heading2_style)
                else:
                    # Check for bold text in runs
                    has_bold = any(run.bold for run in para.runs if run.text.strip())
                    if has_bold and len(text) < 100:
                        # Short bold text might be a heading
                        p = Paragraph(f"<b>{text}</b>", normal_style)
                    else:
                        p = Paragraph(text, normal_style)
                
                story.append(p)
            
            # Add tables if any
            for table in doc.tables:
                story.append(Spacer(1, 0.2 * inch))
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        p = Paragraph(row_text, normal_style)
                        story.append(p)
                story.append(Spacer(1, 0.2 * inch))
            
            # Build PDF
            print(f"  🔨 Building PDF...")
            pdf_doc.build(story)
            
            print(f"  ✅ PDF created: {temp_pdf_path.name}")
            
            # Read PDF
            with open(temp_pdf_path, 'rb') as f:
                pdf_content = f.read()
            
            output = BytesIO(pdf_content)
            output.seek(0)
            
            # Generate final filename
            final_filename = file.filename.rsplit('.', 1)[0] + '.pdf'
            
            print(f"✅ Conversion complete: {final_filename} ({len(pdf_content)} bytes)")
            
            return StreamingResponse(
                output,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename={final_filename}",
                    "X-Conversion-Method": "python-docx-reportlab"
                }
            )
            
        finally:
            # Cleanup temporary files
            print(f"  🗑️ Cleaning up temporary files...")
            for temp_file in [temp_word_path, temp_pdf_path]:
                try:
                    if temp_file.exists():
                        os.remove(temp_file)
                        print(f"    ✓ Removed: {temp_file.name}")
                except Exception as e:
                    print(f"    ✗ Failed to remove {temp_file.name}: {str(e)}")
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Word to PDF error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Conversion failed: {str(e)}")

# ==========================================
# UTILITY ENDPOINTS
# ==========================================

@router.get("/stats")
async def get_stats():
    """Get basic statistics"""
    total_sessions = len(SESSIONS)
    total_files = sum(1 for _ in settings.UPLOAD_DIR.glob("*.pdf"))
    
    return {
        "active_sessions": total_sessions,
        "total_files": total_files
    }


@router.get("/pdfs/{pdf_id}/pages")
async def get_pdf_pages(pdf_id: str):
    """Get all pages of a PDF"""
    session_data = SESSIONS.get(pdf_id)
    if not session_data:
        raise HTTPException(404, "PDF not found")
    
    return {
        "pages": session_data['pages'],
        "total_pages": session_data['total_pages'],
        "filename": session_data['original_filename']
    }


@router.get("/pdfs/{pdf_id}/versions")
async def get_pdf_versions(pdf_id: str):
    """Get version history of a PDF (placeholder)"""
    return {
        "versions": [],
        "message": "Version history not implemented"
    }


@router.delete("/pdfs/{pdf_id}")
async def delete_pdf(pdf_id: str):
    """Delete a PDF and its session"""
    session_data = SESSIONS.get(pdf_id)
    if not session_data:
        raise HTTPException(404, "PDF not found")
    
    try:
        file_path = Path(session_data['file_path'])
        if file_path.exists():
            os.remove(file_path)
        
        for page in session_data['pages']:
            thumb_path = settings.THUMBNAIL_DIR / f"{pdf_id}_page_{page['page_num']}.png"
            if thumb_path.exists():
                os.remove(thumb_path)
        
        del SESSIONS[pdf_id]
        
        return {"message": "PDF deleted successfully"}
    
    except Exception as e:
        print(f"❌ Delete error: {str(e)}")
        raise HTTPException(500, f"Failed to delete PDF: {str(e)}")


@router.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "active_sessions": len(SESSIONS),
        "timestamp": datetime.now().isoformat()
    }